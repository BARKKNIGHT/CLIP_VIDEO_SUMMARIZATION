import os
import uuid
import shutil
import base64
import tempfile
import threading
import time
import cv2
import subprocess
import traceback
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_from_directory,
    jsonify,
    flash,
    after_this_request,
    session,
)
import openai
from fpdf import FPDF
from docx import Document
import smtplib
import ssl
from cs50 import SQL
from werkzeug.security import generate_password_hash, check_password_hash
from email.message import EmailMessage
from functools import wraps
from dotenv import load_dotenv
load_dotenv()

# --- Config ---
UPLOAD_FOLDER = "static/uploads"
FRAMES_FOLDER = "static/frames"
ALLOWED_EXTENSIONS = {"mp4", "avi", "mov", "mkv"}
MAX_CONTENT_LENGTH = 500 * 1024 * 1024  # 500MB

openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "supersecretkey")
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["FRAMES_FOLDER"] = FRAMES_FOLDER
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(FRAMES_FOLDER, exist_ok=True)

# --- Email Config ---
MAIL_SERVER = "mail.privateemail.com"
MAIL_PORT = 587
MAIL_USERNAME = os.getenv("MAIL_USERNAME")  # Set in your environment
MAIL_PASSWORD = os.getenv("MAIL_PASSWORD")  # Set in your environment
MAIL_FROM = MAIL_USERNAME

# --- CS50 SQL DB ---
DB_PATH = os.path.join(os.path.dirname(__file__), 'app.db')
if not os.path.exists(DB_PATH):
    open(DB_PATH, 'a').close()
db = SQL(f"sqlite:///{DB_PATH}")

def init_db():
    db.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT NOT NULL UNIQUE,
            email TEXT NOT NULL UNIQUE,
            password_hash TEXT NOT NULL,
            is_verified INTEGER DEFAULT 0,
            verify_token TEXT,
            reset_token TEXT
        )
    """)
    db.execute("""
        CREATE TABLE IF NOT EXISTS videos (
            id TEXT PRIMARY KEY,
            filename TEXT NOT NULL,
            user_id INTEGER NOT NULL,
            uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    """)

init_db()

# --- Email Sending Helper ---
def send_email(to, subject, body, html_body=None):
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = MAIL_FROM
    msg["To"] = to
    msg.set_content(body)
    if html_body:
        msg.add_alternative(html_body, subtype="html")
    context = ssl.create_default_context()
    with smtplib.SMTP(MAIL_SERVER, MAIL_PORT, timeout=10) as server:
        server.starttls(context=context)
        server.login(MAIL_USERNAME, MAIL_PASSWORD)
        server.send_message(msg)
        print("Sent!")
# --- Auth Routes ---

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        username = request.form.get("username")
        email = request.form.get("email")
        password = request.form.get("password")
        if not username or not email or not password:
            flash("All fields required.", "danger")
            return redirect(url_for("register"))
        if db.execute("SELECT * FROM users WHERE username = ? OR email = ?", username, email):
            flash("Username or email already exists.", "danger")
            return redirect(url_for("register"))
        password_hash = generate_password_hash(password)
        db.execute(
            "INSERT INTO users (username, email, password_hash) VALUES (?, ?, ?)",
            username, email, password_hash
        )
        flash("Registration successful! You can now log in.", "success")
        return redirect(url_for("login"))
    return render_template("register.html")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")
        user = db.execute("SELECT * FROM users WHERE username = ?", username)
        if not user or not check_password_hash(user[0]["password_hash"], password):
            flash("Invalid username or password.", "danger")
            return redirect(url_for("login"))
        session["user_id"] = user[0]["id"]
        session["username"] = user[0]["username"]
        flash("Logged in successfully.", "success")
        return redirect(url_for("index"))
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.clear()
    flash("Logged out.", "info")
    return redirect(url_for("login"))

# --- Protect routes (example) ---
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if "user_id" not in session:
            flash("Please log in.", "warning")
            return redirect(url_for("login"))
        return f(*args, **kwargs)
    return decorated_function

# --- Add session config ---
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_TYPE"] = "filesystem"

# --- Thread-safe Status ---
from threading import Lock

processing_status = {}
status_lock = Lock()

def set_status(video_id, status):
    with status_lock:
        processing_status[video_id] = status

def get_status(video_id):
    with status_lock:
        return processing_status.get(video_id)

# --- Helper Functions ---

def allowed_file(filename):
    """Check if the file has an allowed extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_frames(video_path, frames_dir, num_frames=30):
    """Extract representative frames from the video."""
    cap = cv2.VideoCapture(video_path)
    total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    fps = cap.get(cv2.CAP_PROP_FPS)
    duration = total_frames / fps if fps else 0
    if duration == 0:
        cap.release()
        return [], 0
    # 15 frames per minute, max 60
    minutes = max(1, int(duration // 60))
    num_frames = min(15 * minutes, 60)
    frame_indices = [
        int(i * total_frames / num_frames) for i in range(num_frames)
    ]
    frame_files = []
    for idx, frame_idx in enumerate(frame_indices):
        cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
        ret, frame = cap.read()
        if ret:
            frame_filename = f"frame_{uuid.uuid4().hex[:8]}.jpg"
            frame_path = os.path.join(frames_dir, frame_filename)
            cv2.imwrite(frame_path, frame)
            frame_files.append(frame_filename)
    cap.release()
    return frame_files, duration

def transcribe_audio(audio_path):
    """Transcribe audio using OpenAI Whisper."""
    with open(audio_path, "rb") as audio_file:
        transcript = openai.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file
        )
    return transcript.text

def extract_audio(video_path, audio_path):
    os.makedirs(os.path.dirname(audio_path), exist_ok=True)

    command = [
        "ffmpeg", "-y",
        "-i", video_path,
        "-vn",
        "-acodec", "pcm_s16le",
        "-ar", "16000",
        "-ac", "1",
        audio_path
    ]

    result = subprocess.run(
        command,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE
    )

    if result.returncode != 0:
        stderr = result.stderr.decode(errors="ignore").lower()

        if "audio" in stderr or "stream" in stderr:
            print("[INFO] No audio stream detected. Continuing without transcript.")
            return False

        raise RuntimeError(f"FFmpeg failed:\n{stderr}")

    return True


def describe_frames(frames_dir, frame_files):
    """Describe each frame using OpenAI Vision model."""
    descriptions = []
    for frame_file in frame_files:
        frame_path = os.path.join(frames_dir, frame_file)
        with open(frame_path, "rb") as img_file:
            frame_b64 = base64.b64encode(img_file.read()).decode("utf-8")
        response = openai.chat.completions.create(
            model="gpt-4o",  # Updated to correct model name
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Describe this video frame."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{frame_b64}"
                            },
                        },
                    ],
                }
            ],
            max_tokens=50,
        )
        descriptions.append(response.choices[0].message.content.strip())
    return descriptions

def summarize(descriptions, transcript):
    """Summarize the video using frame descriptions and transcript."""
    prompt = (
        "You are given a video. Here are the frame descriptions:\n"
        + "\n".join(f"Frame {i+1}: {desc}" for i, desc in enumerate(descriptions))
        + "\n\nAnd here is the transcript of the audio:\n"
        + transcript
        + "\n\nSummarize the video content, combining visual and audio information."
    )
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=500,
    )
    return response.choices[0].message.content.strip()

def generate_pdf(transcript, summary, output_path):
    """Generate a PDF file with transcript and summary, supporting Unicode."""
    from fpdf import FPDF
    pdf = FPDF()
    pdf.add_page()
    # Add Unicode font (DejaVuSans)
    font_path = os.path.join('static', 'DejaVuSans.ttf')
    pdf.add_font('DejaVu', '', font_path, uni=True)
    pdf.set_font('DejaVu', '', 12)
    pdf.cell(0, 10, "Transcript", ln=True)
    pdf.multi_cell(0, 10, transcript)
    pdf.cell(0, 10, "Summary", ln=True)
    pdf.multi_cell(0, 10, summary)
    pdf.output(output_path)

def generate_docx(transcript, summary, output_path):
    """Generate a DOCX file with transcript and summary."""
    doc = Document()
    doc.add_heading("Transcript", level=1)
    doc.add_paragraph(transcript)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(summary)
    doc.save(output_path)

# --- Background Processing ---

def process_video(video_id, video_path, frames_dir):
    try:
        set_status(video_id, {"status": "Extracting frames...", "progress": 10})
        frame_files, duration = extract_frames(video_path, frames_dir)
        
        set_status(video_id, {"status": "Extracting audio...", "progress": 30})
        audio_path = os.path.join(frames_dir, f"{video_id}.wav")
        
        has_audio = extract_audio(video_path, audio_path)
        
        if has_audio:
            set_status(video_id, {"status": "Transcribing audio...", "progress": 50})
            transcript = transcribe_audio(audio_path)
        else:
            transcript = ""
        
        set_status(video_id, {"status": "Describing frames...", "progress": 70})
        descriptions = describe_frames(frames_dir, frame_files)
        
        set_status(video_id, {"status": "Summarizing...", "progress": 90})
        summary = summarize(descriptions, transcript)
        
        set_status(video_id, {
            "status": "Completed",
            "progress": 100,
            "frame_files": frame_files,
            "descriptions": descriptions,
            "transcript": transcript,
            "summary": summary,
            "duration": duration,
            "video_path": video_path,
            "frames_dir": frames_dir,
        })
    except Exception as e:
        traceback.print_exc()
        set_status(video_id, {"status": f"Error: {str(e)}", "progress": -1})

# --- Routes ---

@app.route("/", methods=["GET", "POST"])
@login_required
def index():
    user_id = session["user_id"]
    if request.method == "POST":
        if "video" not in request.files:
            flash("No video file part", "danger")
            return redirect(request.url)
        file = request.files["video"]
        if file.filename == "":
            flash("No selected file", "danger")
            return redirect(request.url)
        if file and allowed_file(file.filename):
            video_id = uuid.uuid4().hex
            video_filename = f"{video_id}_{file.filename}"
            video_path = os.path.join(app.config["UPLOAD_FOLDER"], video_filename)
            file.save(video_path)
            frames_dir = os.path.join(app.config["FRAMES_FOLDER"], video_id)
            os.makedirs(frames_dir, exist_ok=True)
            # Save video ownership in DB
            try:
                db.execute(
                    "INSERT INTO videos (id, filename, user_id) VALUES (?, ?, ?)",
                    video_id, video_filename, user_id
                )
            except Exception as e:
                print(f"[ERROR] Failed to insert video into DB: {e}")
                flash("Failed to save video info to database.", "danger")
                return redirect(request.url)
            # Start background processing
            threading.Thread(
                target=process_video, args=(video_id, video_path, frames_dir)
            ).start()
            return redirect(url_for("progress", video_id=video_id))
        else:
            flash("Invalid file type", "danger")
            return redirect(request.url)
    # List only user's processed videos
    videos = []
    user_videos = db.execute("SELECT * FROM videos WHERE user_id = ? ORDER BY uploaded_at DESC", user_id)
    for v in user_videos:
        video_id = v["id"]
        status = get_status(video_id)
        if status and status.get("progress") == 100:
            videos.append({
                "video_id": video_id,
                "filename": v["filename"],
                "duration": status.get("duration", 0),
            })
    return render_template("index.html", videos=videos)

@app.route("/results/<video_id>")
@login_required
def results(video_id):
    user_id = session["user_id"]
    # Check ownership
    video = db.execute("SELECT * FROM videos WHERE id = ? AND user_id = ?", video_id, user_id)
    if not video:
        flash("You do not have access to this video.", "danger")
        return redirect(url_for("index"))
    status = get_status(video_id)
    if not status or status.get("progress") != 100:
        flash("Processing not complete or failed.", "danger")
        return redirect(url_for("index"))
    video_file = video[0]["filename"]
    return render_template(
        "results.html",
        video_id=video_id,
        video_file=video_file,
        frame_descs=zip(status["frame_files"], status["descriptions"]),
        transcript=status["transcript"],
        summary=status["summary"],
        duration=status["duration"],
    )

@app.route("/download/<video_id>/<filetype>")
@login_required
def download_file(video_id, filetype):
    user_id = session["user_id"]
    # Check ownership
    video = db.execute("SELECT * FROM videos WHERE id = ? AND user_id = ?", video_id, user_id)
    if not video:
        flash("You do not have access to this video.", "danger")
        return redirect(url_for("index"))
    status = get_status(video_id)
    if not status or status.get("progress") != 100:
        flash("Processing not complete or failed.", "danger")
        return redirect(url_for("index"))
    transcript = status["transcript"]
    summary = status["summary"]
    temp_dir = tempfile.mkdtemp()
    if filetype == "pdf":
        output_path = os.path.join(temp_dir, f"{video_id}.pdf")
        generate_pdf(transcript, summary, output_path)
        filename = f"{video_id}.pdf"
    elif filetype == "docx":
        output_path = os.path.join(temp_dir, f"{video_id}.docx")
        generate_docx(transcript, summary, output_path)
        filename = f"{video_id}.docx"
    else:
        flash("Invalid file type", "danger")
        shutil.rmtree(temp_dir)
        return redirect(url_for("results", video_id=video_id))
    import gc
    gc.collect()  # Force garbage collection to close file handles
    time.sleep(0.1)  # Small delay to ensure file is closed (Windows)
    @after_this_request
    def cleanup(response):
        try:
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"[WARN] Could not remove temp dir: {e}")
        return response
    return send_from_directory(temp_dir, filename, as_attachment=True)

@app.route("/static/uploads/<filename>")
@login_required
def uploaded_file(filename):
    # Only allow access to user's own videos
    user_id = session["user_id"]
    video = db.execute("SELECT * FROM videos WHERE filename = ? AND user_id = ?", filename, user_id)
    if not video:
        flash("You do not have access to this file.", "danger")
        return redirect(url_for("index"))
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename)

@app.route("/static/frames/<video_id>/<filename>")
@login_required
def frame_file(video_id, filename):
    user_id = session["user_id"]
    video = db.execute("SELECT * FROM videos WHERE id = ? AND user_id = ?", video_id, user_id)
    if not video:
        flash("You do not have access to this file.", "danger")
        return redirect(url_for("index"))
    return send_from_directory(os.path.join(app.config["FRAMES_FOLDER"], video_id), filename)

@app.route("/progress/<video_id>")
@login_required
def progress(video_id):
    return render_template("progress.html", video_id=video_id)

@app.route("/progress_status/<video_id>")
@login_required
def progress_status(video_id):
    status = get_status(video_id) or {"status": "Not found", "progress": -1}
    return jsonify(status)

@app.route("/delete/<video_id>", methods=["POST"])
@login_required
def delete_video(video_id):
    user_id = session["user_id"]
    # Check ownership
    video = db.execute("SELECT * FROM videos WHERE id = ? AND user_id = ?", video_id, user_id)
    if not video:
        flash("You do not have access to this video.", "danger")
        return redirect(url_for("index"))
    # Remove video file
    video_file = video[0]["filename"]
    video_path = os.path.join(app.config["UPLOAD_FOLDER"], video_file)
    if os.path.exists(video_path):
        try:
            os.remove(video_path)
        except Exception as e:
            print(f"[WARN] Could not remove video file: {e}")
    # Remove frames directory
    frames_dir = os.path.join(app.config["FRAMES_FOLDER"], video_id)
    if os.path.exists(frames_dir):
        try:
            shutil.rmtree(frames_dir)
        except Exception as e:
            print(f"[WARN] Could not remove frames dir: {e}")
    # Remove from DB
    db.execute("DELETE FROM videos WHERE id = ? AND user_id = ?", video_id, user_id)
    # Remove from processing_status
    set_status(video_id, None)
    flash("Video deleted successfully.", "success")
    return redirect(url_for("index"))

@app.route("/rename/<video_id>", methods=["POST"])
@login_required
def rename_video(video_id):
    user_id = session["user_id"]
    new_name = request.form.get("new_name", "").strip()
    if not new_name:
        flash("New name cannot be empty.", "danger")
        return redirect(url_for("results", video_id=video_id))
    # Check ownership
    video = db.execute("SELECT * FROM videos WHERE id = ? AND user_id = ?", video_id, user_id)
    if not video:
        flash("You do not have access to this video.", "danger")
        return redirect(url_for("index"))
    old_filename = video[0]["filename"]
    # Extract extension
    ext = old_filename.split(".")[-1]
    # Sanitize new name (remove extension if present)
    if new_name.lower().endswith(f".{ext.lower()}"):
        new_name = new_name[:-(len(ext)+1)]
    # Compose new filename
    new_filename = f"{video_id}_{new_name}.{ext}"
    old_path = os.path.join(app.config["UPLOAD_FOLDER"], old_filename)
    new_path = os.path.join(app.config["UPLOAD_FOLDER"], new_filename)
    # Rename file
    try:
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
    except Exception as e:
        flash(f"Failed to rename file: {e}", "danger")
        return redirect(url_for("results", video_id=video_id))
    # Update DB
    db.execute("UPDATE videos SET filename = ? WHERE id = ? AND user_id = ?", new_filename, video_id, user_id)
    flash("Video renamed successfully.", "success")
    return redirect(url_for("results", video_id=video_id))

# --- Templates ---
# Place index.html, progress.html, results.html in /templates

if __name__ == "__main__":
    app.run(debug=True, port=8000, host="0.0.0.0")
